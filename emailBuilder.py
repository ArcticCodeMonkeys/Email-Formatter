# Online Python - IDE, Editor, Compiler, Interpreter

import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
def remove_default(appliance_list, min_length=1):
    if len(appliance_list) > min_length:
        appliance_list.pop(0)
    return appliance_list


excel_file = filedialog.askopenfilename(title="Select Excel File", filetypes=[("Excel files", "*.xlsx")])
word_template = filedialog.askopenfilename(title="Select Word Template", filetypes=[("Word files", "*.docx")])
if not excel_file or not word_template:
    messagebox.showerror("Input Error", "Please select both an Excel file and a Word template.")
    exit()


df = pd.read_excel(excel_file)

print(str(len(df)) + " Projects Found")

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add underline if it is given
    if underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(u)
    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)


    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

hyperlinks = {
    "FDPC4314AS": ["https://www.frigidaire.ca/Kitchen/Dishwashers/Dishwasher/FDPC4314AS/", "Frigidaire 24\" Built-In Dishwasher"],
    "FFBD1831US": ["https://www.frigidaire.ca/Kitchen/Dishwashers/Dishwasher/FFBD1831US/", "Frigidaire 18\" Built-In Dishwasher"],
    "FRBG1224AV": ["https://www.frigidaire.ca/Kitchen/Refrigerators/Bottom-Freezer/FRBG1224AV/", "Frigidaire 24\" Bottom Freezer Refrigerator"],
    "FFET1222UV": ["https://www.frigidaire.ca/Kitchen/Refrigerators/Top-Freezer-Refrigerators/FFET1222UV/", "Frigidaire 24\" Top Freezer Refrigerator"],
    "GRBN2012AF": ["https://www.frigidaire.ca/Kitchen/Refrigerators/Bottom-Freezer/GRBN2012AF/", "Frigidaire Gallery 30\" Bottom Freezer Refrigerator"],
    "FFHT1835VS": ["https://www.frigidaire.com/en/p/kitchen/refrigerators/top-freezer-refrigerators/FFHT1835VS", "Frigidare 30\" Top Freezer Refrigerator"],
    "FRFG1723AV": ["https://www.frigidaire.ca/Kitchen/Refrigerators/French-Door-Refrigerators/FRFG1723AV/", "Frigidare 33\" Bottom Freezer"],
    "FRFC2323AS": ["https://www.frigidaire.ca/Kitchen/Refrigerators/French-Door-Refrigerators/FRFC2323AS/", "Frigidaire 36\" Bottom Freezer Refrigerator"],
    "FCFE242CAS": ["https://www.frigidaire.ca/Kitchen/Ranges/Electric-Ranges/FCFE242CAS/", "Frigidaire 24\" Free Standing Ceran Top Range"],
    "FCFC241CAW": ["https://www.frigidaire.ca/Kitchen/Ranges/Electric-Ranges/FCFC241CAW/", "Frigidaire 24\" Free Standing Coil Top Range"],
    "FCFE306CAS": ["https://www.frigidaire.ca/Kitchen/Ranges/Electric-Ranges/FCFE306CAS/", "Frigidaire 30\" Front Control Ceran Top Range"],
    "FCFG3062AS": ["https://www.frigidaire.ca/Kitchen/Ranges/Gas-Ranges/FCFG3062AS/", "Frigidaire 30\" Front Control Gas Range"],
    "FCFI308CAS": ["https://www.frigidaire.ca/Kitchen/Ranges/Induction-Ranges/FCFI308CAS/", "Frigidaire 30\" Front Control Induction Range"],
    "PCFI3670AF": ["https://www.frigidaire.ca/Kitchen/Ranges/Induction-Ranges/PCFI3670AF/", "Frigidaire Professional 36\" Free Standing Induction Range"],
    "GCFG3661AF": ["https://www.frigidaire.ca/Kitchen/Ranges/Gas-Ranges/GCFG3661AF/", "Frigidaire Gallery 36\" Gas Range with Air Fry"],
    "ECWS243CAS": ["https://www.electrolux.ca/en/Kitchen-Appliances/Wall-Ovens/Single/ECWS243CAS/", "Electrolux 24\" Single Electric Wall Oven"],
    "FGEW276SPF": ["https://www.frigidaire.ca/Kitchen/Wall-Ovens/Single-Wall-Ovens/FGEW276SPF/", "Frigidaire Gallery 27\" Single Electric Wall Oven"],
    "FCWS3027AS": ["https://www.frigidaire.ca/Kitchen/Wall-Ovens/Single-Wall-Ovens/FCWS3027AS/", "Frigidaire 30\" Single Electric Wall Oven with Fan Convection"],
    "ECCE242CAS": ["https://www.electrolux.ca/en/Kitchen-Appliances/Cooktops/Electric/ECCE242CAS/", "Electrolux 24\" Electric Cooktop"],
    "FFEC3025US": ["https://www.frigidaire.ca/Kitchen/Cooktops/Electric-Cooktops/FFEC3025US/", "Frigidaire 30\" Electric Ceran Cooktop"],
    "GCCG3048AS": ["https://www.frigidaire.ca/Kitchen/Cooktops/Gas-Cooktops/GCCG3048AS/", "Frigidaire Gallery 30\" Gas Cooktop"],
    "FCCI3027AB": ["https://www.frigidaire.ca/Kitchen/Cooktops/Induction-Cooktops/FCCI3027AB/", "Frigidaire 30\" Induction Cooktop"],
    "GCCE3670US": ["https://www.frigidaire.com/en/p/kitchen/cooktops/electric-cooktops/GCCE3670AS", "Frigidaire Gallery 36\" Electric Ceran Cooktop"],
    "GCCG3648AS": ["https://www.frigidaire.com/en/p/kitchen/cooktops/gas-cooktops/GCCG3648AS", "Frigidaire Gallery 36\" Gas Cooktop"],
    "GCCI3667AB": ["https://www.frigidaire.com/en/p/kitchen/cooktops/induction-cooktops/GCCI3667AB", "Frigidaire Gallery 36\" Induction Cooktop"],
    "FCVW3052AS": ["https://www.frigidaire.com/en/p/kitchen/ventilation/FCVW3052AS", "Frigidaire 30\" Undercabinet Ventilation Hood"],
    "UMV1422US":  ["https://www.frigidaire.com/en/p/kitchen/microwaves/over-the-range-microwave/UMV1422US",  "Frigidaire 24\" Over-The-Range Microwave"],
    "FMOS1846BS": ["https://www.frigidaire.com/en/p/kitchen/microwaves/over-the-range-microwave/FMOS1846BS", "Frigidaire 30\" Over-The-Range Microwave"],
    "ELFW7337AW": ["https://www.electrolux.ca/en/Laundry-Appliances/Washing-Machines/Washers/ELFW7337AW/", "Electrolux 27\" Front Load Stackable Washer"],
    "ELFE733CAW": ["https://www.electrolux.ca/en/Laundry-Appliances/Washing-Machines/Dryers/ELFE733CAW/", "Electrolux 27\" Front Load Stackable Electric Dryer"],
    "FLCE752CAW": ["https://www.frigidaire.ca/Washers-Dryers/Laundry-Center/Laundry-Centers/FLCE752CAW/", "Frigidaire 27\" Electric Washer/Dryer Laundry Center"],
    "ELFW4222AW": ["https://www.electrolux.ca/en/Laundry-Appliances/Washing-Machines/Washers/ELFW4222AW/", "Electrolux 24\" Front Load Washer with LuxCare Wash System"],
    "ELFE422CAW": ["https://www.electrolux.ca/en/Laundry-Appliances/Washing-Machines/Dryers/ELFE422CAW/", "Electrolux 24\" Ventless Front Load Dryer"]
}


for index, row in df.iterrows():
    doc = Document(word_template)
    
    dev = f"{row['Developer']}"
    arch = f"{row['Architect']}"
    proj = f"{row['Project Name']}"
    proj = proj.replace("/", "_")
    proj = proj.replace(":", "-")
    fridge = f"{row['Fridge']}"
    fridgeSize = f"{row['Fridge Size']}"
    stove = f"{row['Stove']}"
    stoveSize = f"{row['Stove Top Size']}"
    dishwasherSize = f"{row['Dishwasher Size']}"
    hoodSize = f"{row['Hood Size']}"
    ovenSize = f"{row['Oven Size']}"
    rangeSize = f"{row['Range Size']}"
    microwaveSize = f"{row['Microwave Size']}"
    washerDryer = f"{row['Washer/Dryer']}"
    fridgeModel = []
    stoveModel = ["ECCE242CAS"]
    ovenModel = ["ECWS243CAS"]
    rangeModel = ["FCFE306CAS"]
    dishwasherModel = ["FDPC4314AS"]
    hoodModel = ["FCVW3052AS"]
    microwaveModel = ["UMV1422US"]
    washerModel = ["ELFW7337AW"]
    dryerModel = ["ELFE733CAW"]
    appliance_models = {
    'Stove': stoveModel,
    'Oven': ovenModel,
    'Range': rangeModel,
    'Dishwasher': dishwasherModel,
    'Hood': hoodModel,
    'Microwave': microwaveModel,
    'Washer': washerModel,
    'Dryer': dryerModel
    }
    #REFRIGERATORS
    if(fridge == "Bottom Freezer"):
        #Default Bottom
        fridgeModel.append("FRBG1224AV")
        if("30" in fridgeSize):
            fridgeModel.append("GRBN2012AF")
        if("33" in fridgeSize):
            fridgeModel.append("FRFG1723AV")
        if("36" in fridgeSize):
            fridgeModel.append("FRFC2323AS")
    if(fridge == "Top Freezer"):
        #Default Top
        fridgeModel.append("FFET1222UV")
        if("24" in fridgeSize):
            fridgeModel.append("FFET1222UV")
        if("30" in fridgeSize):
            fridgeModel.append("FFHT1835VS")
    
    #STOVE TOPS
    if("24" in stoveSize):
        stoveModel.append("ECCE242CAS")
    if("30" in stoveSize):
        if(stove == "Ceran Top"):
            stoveModel.append("FFEC3025UB")
        if("Gas" in stove):
            stoveModel.append("GCCG3048AS")
        if(stove == "Induction"):
            stoveModel.append("FCCI3027AB")
    if("36" in stoveSize):
        if(stove == "Ceran Top"):
            stoveModel.append("GCCE3670US")
        if("Gas" in stove):
            stoveModel.append("GCCG3648AS")
        if(stove == "Induction"):
            stoveModel.append("GCCI3667AB")
    
    #WASHERS AND DRYERS
    if("Stacked" in washerDryer):
        washerModel.append("FLCE752CAW")
        dryerModel[0] = "N/A"
    if("Stackable" in washerDryer):
        if("24" in washerDryer):
            washerModel.append("ELFW4222AW")
            dryerModel.append("ELFE422CAW")
        if("27" in washerDryer):
            washerModel.append("ELFW7337AW")
            dryerModel.append("ELFE733CAW")
    
    #OVENS 
    if("24" in ovenSize):
        ovenModel.append("ECWS243CAS")
    if("27" in ovenSize):
        ovenModel.append("FGEW276SPF")
    if("30" in ovenSize):
        ovenModel.append("FCWS3027AS")
    #HOODS 
    if("30" in hoodSize):
        hoodModel.append("FCVW3052AS")
    
    #RANGES 
    if("24" in rangeSize):
        if(stove == "Ceran Top"):
            rangeModel.append("FCFE242CAS")
        if(stove == "Coil Top"):
            rangeModel.append("FCFC241CAW")
    if("30" in rangeSize):
        if(stove == "Ceran Top"):
            rangeModel.append("FCFE306CAS")
        if("Gas" in stove):
            rangeModel.append("FCFG3062AS")
        if(stove == "Induction"):
            rangeModel.append("FCFI308CAS")
    if("36" in rangeSize):
        if("Gas" in stove):
            rangeModel.append("GCFG3661AF")
        if(stove == "Induction"):
            rangeModel.append("PCFI3670AF")
    
    #DISHWASHERS 
    if("18" in dishwasherSize):
        dishwasherModel.append("FFBD1831US")
    if("24" in dishwasherSize):
        dishwasherModel.append("FDPC4314AS")
    #MICROWAVES 
    if("24" in microwaveSize):
        microwaveModel.append("UMV1422US")
    if("30" in microwaveSize):
        microwaveModel.append("FMOS1846BS")
    
    stoveModel = remove_default(stoveModel)
    ovenModel = remove_default(ovenModel)
    rangeModel = remove_default(rangeModel)
    dishwasherModel = remove_default(dishwasherModel)
    hoodModel = remove_default(hoodModel)
    microwaveModel = remove_default(microwaveModel)
    washerModel = remove_default(washerModel)
    dryerModel = remove_default(dryerModel)
    if "N/A" in fridgeModel:
        fridgeModel.remove("N/A")
    if "N/A" in stoveModel:
        stoveModel.remove("N/A")
    if "N/A" in ovenModel:
        ovenModel.remove("N/A")
    if "N/A" in rangeModel:
        rangeModel.remove("N/A")
    if "N/A" in hoodModel:
        hoodModel.remove("N/A")
    if "N/A" in dishwasherModel:
        dishwasherModel.remove("N/A")
    if "N/A" in microwaveModel:
        microwaveModel.remove("N/A")
    if "N/A" in washerModel:
        washerModel.remove("N/A")
    if "N/A" in dryerModel:
        dryerModel.remove("N/A")

    #print(stoveModel + " "  + ovenModel + " " + rangeModel + "" + dishwasherModel + "" + hoodModel + "" + microwaveModel + "" + washerDryerModel)
    #Reformatting the Email and Inserting the Appliances
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for paragraph in doc.paragraphs:
        if "[Project Name]" in paragraph.text:
            paragraph.text = paragraph.text.replace("[Project Name]", proj)
        
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    table = doc.tables[0]
    rowCount = 1
    fridgeCol = table.columns[0]
    cookingCol = table.columns[1]
    dishwasherCol = table.columns[2]
    ventilationCol = table.columns[3]
    laundryCol = table.columns[4]
    
    #Fridge Col
    fridgeCount = 1
    
    for model in fridgeModel:
        if(fridgeCount > rowCount):
            table.add_row()
            rowCount += 1
            for cell in table.rows[rowCount].cells:
                cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
        fridgeCell = fridgeCol.cells[fridgeCount]
        fridgeP = fridgeCell.paragraphs[0]
        fridgeP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        fridgeR = fridgeP.add_run()
        fridgeR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(fridgeP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        fridgeCount += 1
    for cell in fridgeCol.cells:
        cell.width = 3.3
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    
    #Cooking Col
    cookingCount = 1
    for model in stoveModel:
        if(cookingCount > rowCount):
            table.add_row()
            rowCount += 1
        stoveCell = cookingCol.cells[cookingCount]
        stoveP = stoveCell.paragraphs[0]
        stoveP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        stoveR = stoveP.add_run()
        stoveR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(stoveP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        cookingCount += 1
    
    for model in ovenModel:
        if(cookingCount > rowCount):
            table.add_row()
            rowCount += 1
        ovenCell = cookingCol.cells[cookingCount]
        ovenP = ovenCell.paragraphs[0]
        ovenP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        ovenR = ovenP.add_run()
        ovenR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(ovenP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        cookingCount += 1
    for model in rangeModel:
        if(cookingCount > rowCount):
            table.add_row()
            rowCount += 1
            for cell in table.rows[rowCount].cells:
                cell.vertical_alignment= WD_ALIGN_VERTICAL.CENTER
        rangeCell = cookingCol.cells[cookingCount]
        rangeP = rangeCell.paragraphs[0]
        rangeP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rangeR = rangeP.add_run()
        rangeR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(rangeP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        cookingCount += 1
    for cell in cookingCol.cells:
        cell.width = 3.3
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER


    #Ventilation Col
    ventCount = 1
    for model in hoodModel:
        if(ventCount > rowCount):
            table.add_row()
            rowCount += 1
        hoodCell = ventilationCol.cells[ventCount]
        hoodP = hoodCell.paragraphs[0]
        hoodP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        hoodR = hoodP.add_run()
        hoodR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(hoodP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        ventCount += 1
    for model in microwaveModel:
        if(ventCount > rowCount):
            table.add_row()
            rowCount += 1
        microwaveCell = ventilationCol.cells[ventCount]
        microwaveP = microwaveCell.paragraphs[0]
        microwaveP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        microwaveR = microwaveP.add_run()
        microwaveR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(microwaveP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        ventCount += 1
    for cell in ventilationCol.cells:
        cell.width = 3.3
    #Dishwasher Col
    dishwasherCount = 1
    for model in dishwasherModel:
        if(dishwasherCount > rowCount):
            table.add_row()
            rowCount += 1
        dishwasherCell = dishwasherCol.cells[dishwasherCount]
        dishwasherP = dishwasherCell.paragraphs[0]
        dishwasherP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        dishwasherR = dishwasherP.add_run()
        dishwasherR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(dishwasherP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        dishwasherCount += 1
    for cell in dishwasherCol.cells:
        cell.width = 3.3
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    #Laundry Col
    laundryCount = 1
    for model in washerModel:
        if(laundryCount > rowCount):
            table.add_row()
            rowCount += 1
        washerCell = laundryCol.cells[laundryCount]
        washerP = washerCell.paragraphs[0]
        washerP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        washerR = washerP.add_run()
        washerR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(washerP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        laundryCount += 1
    for model in dryerModel:
        if(laundryCount > rowCount):
            table.add_row()
            rowCount += 1
        dryerCell = laundryCol.cells[laundryCount]
        dryerP = dryerCell.paragraphs[0]
        dryerP.alignment=WD_ALIGN_PARAGRAPH.CENTER
        dryerR = dryerP.add_run()
        dryerR.add_picture(f"{model}.png", width=Inches(1))
        add_hyperlink(dryerP, hyperlinks[model][0], hyperlinks[model][1], '0000EE', True)
        laundryCount += 1
    for cell in laundryCol.cells:
        cell.width = 3.3
        cell.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in table.columns:
        for cell in col.cells:
            cell.vertical_alignment= WD_ALIGN_VERTICAL.BOTTOM
    #Saving
    subfolder = 'Mail'
    doc_name = f"{subfolder}/{dev}_{proj}.docx"
    doc.save(doc_name)
    print("Document \"" + doc_name + "\" has been saved to " + subfolder + ".")
    if(arch != "N/A"):
        doc_name = f"{subfolder}/{arch}_{proj}.docx"
        doc.save(doc_name)
        print("Document \"" + doc_name + "\" has been saved to " + subfolder + ".")

print("Finished saving all documents, check the " + subfolder + " folder.")


